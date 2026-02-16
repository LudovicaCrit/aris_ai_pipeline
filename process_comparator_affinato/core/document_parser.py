"""Parser for ARIS Process Documentation (.doc/.docx/.txt)"""

import re
import subprocess
import tempfile
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, Tuple


@dataclass
class Activity:
    """Single process activity."""
    code: str
    title: str
    description: str
    executor: str
    it_system: Optional[str] = None
    control_type: Optional[str] = None
    
    def is_manual(self) -> bool:
        if self.it_system and self.it_system != "-":
            return False
        if self.control_type and "auto" in self.control_type.lower():
            return False
        return True
    
    def __hash__(self):
        return hash(self.code)


@dataclass
class ProcessDocument:
    """Complete ARIS process document."""
    process_name: str
    process_code: str
    macroprocess: str
    owner: str
    activities: list[Activity] = field(default_factory=list)
    diagram_image: Optional[bytes] = None  # PNG bytes of the flow diagram
    diagram_path: Optional[Path] = None    # Path to extracted diagram file
    has_track_changes: bool = False        # Whether document had track changes
    
    def get_executors(self) -> set[str]:
        return {a.executor for a in self.activities if a.executor}
    
    def count_manual_activities(self) -> int:
        return sum(1 for a in self.activities if a.is_manual())
    
    def has_diagram(self) -> bool:
        return self.diagram_image is not None or self.diagram_path is not None


def extract_track_changes_versions(doc_path: Path) -> Tuple[Optional[str], Optional[str], bool]:
    """
    Extract As-Is (without changes) and To-Be (with changes accepted) text from a document with Track Changes.
    
    Returns:
        Tuple of (as_is_text, to_be_text, has_changes)
        - as_is_text: Text with all insertions removed (original version)
        - to_be_text: Text with all changes accepted (final version)
        - has_changes: Whether the document had track changes
    """
    from docx import Document
    from docx.oxml.ns import qn
    
    # Convert to docx if needed
    doc_path = Path(doc_path)
    if doc_path.suffix.lower() == '.doc':
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                str(doc_path), '--outdir', tmpdir
            ], capture_output=True, timeout=60)
            docx_path = Path(tmpdir) / (doc_path.stem + '.docx')
            if not docx_path.exists():
                return None, None, False
            doc = Document(str(docx_path))
    else:
        doc = Document(str(doc_path))
    
    body = doc.element.body
    
    # Check for track changes
    insertions = body.findall('.//' + qn('w:ins'))
    deletions = body.findall('.//' + qn('w:del'))
    
    has_changes = len(insertions) > 0 or len(deletions) > 0
    
    if not has_changes:
        return None, None, False
    
    # Extract text for both versions
    def get_text_recursive(element, include_insertions=True, include_deletions=False):
        """Recursively extract text from element."""
        text_parts = []
        
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag == 'ins':
                # Insertion - include only in To-Be (include_insertions=True)
                if include_insertions:
                    text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
            elif tag == 'del':
                # Deletion - include only in As-Is (include_deletions=True)
                if include_deletions:
                    text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
            elif tag == 't':
                # Text node
                if child.text:
                    text_parts.append(child.text)
            elif tag == 'tab':
                text_parts.append('\t')
            elif tag == 'br':
                text_parts.append('\n')
            elif tag == 'p':
                # Paragraph - add newline after
                text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
                text_parts.append('\n')
            elif tag == 'tr':
                # Table row - add newline after
                text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
                text_parts.append('\n')
            elif tag == 'tc':
                # Table cell - add tab after
                text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
                text_parts.append('\t')
            else:
                # Recurse into other elements
                text_parts.append(get_text_recursive(child, include_insertions, include_deletions))
        
        return ''.join(text_parts)
    
    # As-Is: exclude insertions, include deletions
    as_is_text = get_text_recursive(body, include_insertions=False, include_deletions=True)
    
    # To-Be: include insertions, exclude deletions
    to_be_text = get_text_recursive(body, include_insertions=True, include_deletions=False)
    
    return as_is_text, to_be_text, True


class ARISDocumentParser:
    """Parser for ARIS-exported process documentation."""
    
    FIELD_PATTERNS = {
        # Patterns handle both newline-separated and compact formats
        'title': re.compile(r'TITOLO\s*\n?(.+?)(?=\n?(?:DESCRIZIONE|ALTRO|ESECUTORE|$))', re.DOTALL),
        'description': re.compile(r'DESCRIZIONE\s*\n?(.+?)(?=\n?(?:GESTIONE ANOMALIA|ALTRO STRUMENTO|ESECUTORE|TIPO|$))', re.DOTALL),
        'executor': re.compile(r'ESECUTORE\s*\n?(.+?)(?=\n?(?:APPLICATIVO|ALTRO|TIPO|$))', re.DOTALL),
        'it_system': re.compile(r'APPLICATIVO INFORMATICO\s*\n?(.+?)(?=\n?(?:EFFICACIA|NATURA|MODALITA|$|\d{3}))', re.DOTALL),
        'control_type': re.compile(r'\n?TIPO\s*\n?(.+?)(?=\n?(?:SCOPO|$))', re.DOTALL),
    }
    
    def __init__(self):
        # antiword check moved to extract_text, only when .doc file is processed
        pass
    
    def _check_antiword(self):
        """Check if antiword is available (only called for .doc files)."""
        try:
            subprocess.run(['antiword', '--version'], capture_output=True, check=False)
        except FileNotFoundError:
            raise RuntimeError("antiword not found. Install with: apt-get install antiword (Linux) or use .txt/.docx files")
    
    def _convert_doc_to_docx(self, doc_path: Path, output_dir: Path) -> Optional[Path]:
        """Convert .doc to .docx using LibreOffice."""
        try:
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'docx',
                str(doc_path), '--outdir', str(output_dir)
            ], capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                docx_path = output_dir / (doc_path.stem + '.docx')
                if docx_path.exists():
                    return docx_path
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        return None
    
    def _convert_metafile_to_png(self, metafile_path: Path, output_dir: Path) -> Optional[Path]:
        """Convert EMF/WMF to PNG using LibreOffice."""
        try:
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'png',
                str(metafile_path), '--outdir', str(output_dir)
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                png_path = output_dir / (metafile_path.stem + '.png')
                if png_path.exists():
                    return png_path
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        return None
    
    def extract_diagram(self, file_path: Path, output_dir: Optional[Path] = None) -> Optional[Path]:
        """
        Extract flow diagram image from Word document.
        
        Returns path to PNG file, or None if no diagram found.
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if output_dir is None:
            output_dir = Path(tempfile.mkdtemp(prefix='aris_diagrams_'))
        else:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
        
        docx_path = None
        temp_docx = False
        
        # For .doc files, convert to .docx first
        if suffix == '.doc':
            docx_path = self._convert_doc_to_docx(file_path, output_dir)
            temp_docx = True
            if not docx_path:
                return None
        elif suffix == '.docx':
            docx_path = file_path
        else:
            # .txt files don't have embedded images
            return None
        
        # Extract images from docx using python-docx
        try:
            from docx import Document
            doc = Document(str(docx_path))
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    ext = rel.target_ref.split('.')[-1].lower()
                    
                    # Save the image
                    img_filename = f"{file_path.stem}_diagram.{ext}"
                    img_path = output_dir / img_filename
                    
                    with open(img_path, 'wb') as f:
                        f.write(image_data)
                    
                    # Convert EMF/WMF to PNG
                    if ext in ('emf', 'wmf'):
                        png_path = self._convert_metafile_to_png(img_path, output_dir)
                        if png_path:
                            # Remove original metafile
                            img_path.unlink()
                            return png_path
                    elif ext == 'png':
                        return img_path
                    
                    # For other formats, try to convert
                    return img_path
                    
        except ImportError:
            # python-docx not installed
            pass
        except Exception as e:
            print(f"Warning: Could not extract diagram: {e}")
        finally:
            # Clean up temp docx if created
            if temp_docx and docx_path and docx_path.exists():
                try:
                    docx_path.unlink()
                except:
                    pass
        
        return None
    
    def extract_text(self, file_path: Path) -> str:
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()
        
        if suffix == '.doc':
            # First try antiword
            try:
                self._check_antiword()
                result = subprocess.run(['antiword', str(file_path)], capture_output=True, text=True)
                if result.returncode == 0:
                    text = result.stdout
                    # Check if antiword produced pipe-table format (problematic)
                    # If so, fall back to libreoffice conversion
                    if '|' in text[:2000] and text.count('|') > 20:
                        # antiword produced table format, try libreoffice instead
                        pass
                    else:
                        return text
            except:
                pass
            
            # Fallback: convert to txt with libreoffice
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'txt:Text',
                    str(file_path), '--outdir', tmpdir
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    txt_path = Path(tmpdir) / (file_path.stem + '.txt')
                    if txt_path.exists():
                        return txt_path.read_text(encoding='utf-8', errors='replace')
            
            raise RuntimeError(f"Failed to extract text from {file_path}")
            
        elif suffix == '.docx':
            result = subprocess.run(['pandoc', '-f', 'docx', '-t', 'plain', str(file_path)], capture_output=True, text=True)
            if result.returncode != 0:
                raise RuntimeError(f"pandoc failed: {result.stderr}")
            return result.stdout
        elif suffix == '.txt':
            return file_path.read_text(encoding='utf-8')
        else:
            raise ValueError(f"Unsupported format: {suffix}")
    
    def parse(self, file_path: Path, extract_images: bool = True) -> ProcessDocument:
        file_path = Path(file_path)
        raw_text = self.extract_text(file_path)
        
        doc = self._parse_from_text(raw_text, file_path)
        
        # Extract diagram if requested
        if extract_images:
            diagram_path = self.extract_diagram(file_path)
            if diagram_path and diagram_path.exists():
                doc.diagram_path = diagram_path
                with open(diagram_path, 'rb') as f:
                    doc.diagram_image = f.read()
        
        return doc
    
    def _parse_from_text(self, raw_text: str, file_path: Path = None) -> ProcessDocument:
        """Parse a ProcessDocument from raw text content."""
        process_name = self._extract_process_name(raw_text)
        
        # Fallback to filename if extraction failed and file_path provided
        if process_name == "Unknown Process" and file_path:
            stem = file_path.stem
            stem = re.sub(r'^[\d\._]+\s*', '', stem)
            stem = re.sub(r'_?(as[_-]?is|to[_-]?be|scenario\d*).*$', '', stem, flags=re.IGNORECASE)
            stem = stem.replace('_', ' ').strip()
            if stem:
                process_name = stem
        
        doc = ProcessDocument(
            process_name=process_name,
            process_code=self._extract_process_code(raw_text),
            macroprocess=self._extract_field(raw_text, r'Macroprocesso\s*\n(.+?)\n'),
            owner=self._extract_field(raw_text, r'Owner\s*\n(.+?)\n'),
        )
        doc.activities = self._extract_activities(raw_text)
        
        return doc
        
        # Extract diagram if requested
        if extract_images:
            diagram_path = self.extract_diagram(file_path)
            if diagram_path and diagram_path.exists():
                doc.diagram_path = diagram_path
                with open(diagram_path, 'rb') as f:
                    doc.diagram_image = f.read()
        
        return doc
    
    def _extract_process_name(self, text: str) -> str:
        # Try multiple patterns to handle different formats
        patterns = [
            # Format: "Flusso di processo\n\nNAME\n"
            r'Flusso di processo\s*\n+(.+?)\n',
            # Format: "Flusso di processo    NAME"
            r'Flusso di processo\s+(\d+[\.\d]*\s*[^\n]+)',
            # Format from tables: "Flusso di processo" followed by code and name
            r'Flusso di\s*processo\s*\n*\s*(\d+[\.\d]*[A-Za-z]*\s+[^\n]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Clean up the name
                name = re.sub(r'\s+', ' ', name)  # normalize whitespace
                if name and name != '-' and len(name) > 3:
                    return name
        
        # Fallback: try to extract from filename pattern if text extraction failed
        return "Unknown Process"
    
    def _extract_process_code(self, text: str) -> str:
        match = re.search(r'(\d+\.\d+\.?\d*\.?[A-Z]*)\s', text)
        return match.group(1) if match else ""
    
    def _extract_field(self, text: str, pattern: str) -> str:
        match = re.search(pattern, text)
        return match.group(1).strip() if match else ""
    
    def _extract_activities(self, text: str) -> list[Activity]:
        activities = []
        
        # Normalize: replace \x07 (bell/tab) with newline for consistent parsing
        text_normalized = text.replace('\x07', '\n')
        
        # Strategy 1: Original format with newlines (\n010\n)
        parts = re.split(r'\n(\d{3})\n', text_normalized)
        if len(parts) > 2:
            i = 1
            while i < len(parts) - 1:
                code = parts[i]
                content = parts[i + 1] if i + 1 < len(parts) else ""
                
                section_header = ""
                if not content.strip().startswith("TITOLO"):
                    lines = content.split("\n", 2)
                    if len(lines) > 1:
                        section_header = lines[0].strip()
                        content = "\n".join(lines[1:])
                
                activity = self._parse_activity_block(code, content, section_header)
                if activity:
                    activities.append(activity)
                i += 2
        
        # Strategy 2: Compact format (010TITOLO...020TITOLO...)
        if not activities:
            pattern = r'(\d{3})TITOLO\s*\n?(.+?)(?=\d{3}TITOLO|$)'
            matches = re.findall(pattern, text_normalized, re.DOTALL)
            
            for code, content in matches:
                content = "TITOLO\n" + content
                activity = self._parse_activity_block(code, content)
                if activity:
                    activities.append(activity)
        
        # Strategy 3: Pipe-table format (|010|TITOLO|...|)
        if not activities:
            # Find all activity codes in pipe format
            # Pattern matches |CODE| or |CODE |
            code_pattern = r'\|(\d{3})\s*\|'
            code_matches = list(re.finditer(code_pattern, text))
            
            for i, match in enumerate(code_matches):
                code = match.group(1)
                start = match.end()
                
                # Find end (next code or end of reasonable content)
                if i + 1 < len(code_matches):
                    end = code_matches[i + 1].start()
                else:
                    end = min(start + 3000, len(text))
                
                block = text[start:end]
                
                # Clean up pipe-table format: remove pipes and normalize
                block_clean = re.sub(r'\|', '\n', block)
                block_clean = re.sub(r'\n\s*\n', '\n', block_clean)
                block_clean = re.sub(r'^\s+', '', block_clean, flags=re.MULTILINE)
                
                # Check for section header (like "1A - Richieste riservate Direzione")
                # It appears before TITOLO in the cleaned block
                section_header = ""
                lines = block_clean.strip().split('\n')
                
                # Find where TITOLO is
                titolo_idx = -1
                for idx, line in enumerate(lines):
                    if 'TITOLO' in line.upper():
                        titolo_idx = idx
                        break
                
                # If there's content before TITOLO, it might be a section header
                if titolo_idx > 0:
                    potential_header = lines[0].strip()
                    # Section headers typically start with pattern like "1A -" or "1B -"
                    if re.match(r'^\d+[A-Z]?\s*-', potential_header) or (potential_header and 'TITOLO' not in potential_header.upper() and len(potential_header) > 3):
                        section_header = potential_header
                        # Remove section header from block for parsing
                        block_clean = '\n'.join(lines[1:])
                
                if 'TITOLO' in block_clean.upper():
                    activity = self._parse_activity_block(code, block_clean, section_header)
                    if activity:
                        activities.append(activity)
        
        # Strategy 4: Look for TITOLO blocks and find codes nearby
        if not activities:
            pattern = r'(\d{3})\s*\n?TITOLO\s*\n(.+?)(?=\d{3}\s*\n?TITOLO|LEGENDA|$)'
            matches = re.findall(pattern, text_normalized, re.DOTALL | re.IGNORECASE)
            
            for code, content in matches:
                content = "TITOLO\n" + content
                activity = self._parse_activity_block(code, content)
                if activity:
                    activities.append(activity)
        
        return activities
    
    def _parse_activity_block(self, code: str, content: str, section_header: str = "") -> Optional[Activity]:
        def extract(key: str) -> str:
            pattern = self.FIELD_PATTERNS.get(key)
            if pattern:
                match = pattern.search(content)
                if match:
                    return match.group(1).strip()
            return ""
        
        title = extract('title')
        if not title:
            return None
        
        if section_header:
            title = f"{section_header} - {title}"
        
        return Activity(
            code=code,
            title=title,
            description=extract('description'),
            executor=extract('executor'),
            it_system=extract('it_system') or "-",
            control_type=extract('control_type'),
        )


def parse_aris_document(file_path: str | Path) -> ProcessDocument:
    parser = ARISDocumentParser()
    return parser.parse(Path(file_path))


def parse_document_with_track_changes(file_path: str | Path) -> Tuple[ProcessDocument, ProcessDocument, bool]:
    """
    Parse a document that may contain Track Changes.
    
    If Track Changes are detected:
        - Returns (as_is_doc, to_be_doc, True)
        - as_is_doc: Document with changes rejected (original)
        - to_be_doc: Document with changes accepted (modified)
    
    If no Track Changes:
        - Returns (doc, doc, False)
        - Same document returned twice
    
    This allows a single file with Track Changes to serve as both As-Is and To-Be.
    """
    file_path = Path(file_path)
    parser = ARISDocumentParser()
    
    # Try to extract track changes versions
    as_is_text, to_be_text, has_changes = extract_track_changes_versions(file_path)
    
    if has_changes and as_is_text and to_be_text:
        # Parse both versions from the extracted text
        as_is_doc = parser._parse_from_text(as_is_text, file_path)
        as_is_doc.has_track_changes = True
        
        to_be_doc = parser._parse_from_text(to_be_text, file_path)
        to_be_doc.has_track_changes = True
        
        # Extract diagram from the original file (same for both versions)
        diagram_path = parser.extract_diagram(file_path)
        if diagram_path and diagram_path.exists():
            with open(diagram_path, 'rb') as f:
                diagram_bytes = f.read()
            as_is_doc.diagram_path = diagram_path
            as_is_doc.diagram_image = diagram_bytes
            to_be_doc.diagram_path = diagram_path
            to_be_doc.diagram_image = diagram_bytes
        
        return as_is_doc, to_be_doc, True
    else:
        # No track changes, parse normally
        doc = parser.parse(file_path)
        return doc, doc, False