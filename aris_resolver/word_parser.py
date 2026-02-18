"""
Parser per i Word dei Process Owner.
Estrae entità (attività, esecutori, applicativi, controlli) dal testo
e preserva le relazioni tra di loro.

Ogni attività porta con sé il nome dell'esecutore e dell'applicativo
estratti dallo stesso blocco tabellare. Questo permette al Resolver
di confrontare le connessioni Word (to-be) con quelle ARIS (as-is).

Supporta:
- File .doc (RTF) tramite striprtf
- File .doc (binario OLE2) tramite antiword
- File .docx tramite python-docx
"""

import re
import subprocess
from models import WordEntity


def read_word_file(filepath: str) -> str:
    """
    Legge un file Word e restituisce il testo.

    Per i .doc prova prima striprtf (file RTF).
    Se fallisce (file binario OLE2), usa antiword.
    """
    if filepath.endswith('.doc'):
        # Prova prima come RTF
        try:
            from striprtf.striprtf import rtf_to_text
            with open(filepath, 'r', encoding='cp1252', errors='replace') as f:
                raw = f.read()
            text = rtf_to_text(raw)
            # Verifica che abbia estratto qualcosa di utile
            if 'TITOLO' in text or 'DESCRIZIONE' in text:
                return text
        except Exception:
            pass

        # Fallback: antiword per .doc binari (OLE2)
        try:
            result = subprocess.run(
                ['antiword', filepath],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        raise ValueError(
            f"Impossibile leggere {filepath}: "
            "né striprtf né antiword hanno funzionato. "
            "Installa antiword con: sudo apt install antiword"
        )

    elif filepath.endswith('.docx'):
        import docx
        doc = docx.Document(filepath)
        # Estrai testo da paragraphs
        text = "\n".join(p.text for p in doc.paragraphs)
        # Se non contiene TITOLO, estrai dalle tabelle
        # ricostruendo il formato atteso: \n010\nTITOLO\n...
        if 'TITOLO' not in text:
            parts = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    # Formato ARIS: col 0 = codice attività, col 1 = contenuto
                    #               col 2 = codice controllo, col 3 = contenuto controllo
                    # Ricostruisci: codice + contenuto per attività e controlli
                    if len(cells) >= 2:
                        # Colonne 0-1: attività
                        if cells[0] and cells[1] and 'TITOLO' in cells[1]:
                            parts.append(f"\n{cells[0]}\x07TITOLO\n{cells[1].split('TITOLO', 1)[1].lstrip()}")
                        # Colonne 2-3: controlli
                        if len(cells) >= 4 and cells[2] and cells[3] and 'TITOLO' in cells[3]:
                            parts.append(f"\n{cells[2]}\x07TITOLO\n{cells[3].split('TITOLO', 1)[1].lstrip()}")
            text = "\n".join(parts)
        return text

    else:
        raise ValueError(f"Formato non supportato: {filepath}")


def extract_entities(text: str) -> list[WordEntity]:
    """
    Estrae entità dal testo del Word.

    Parsa la struttura tabellare dei Word di Reale Mutua:
    - Formato RTF (striprtf): '010|TITOLO...'
    - Formato antiword: '010TITOLO...'
    - Campi ESECUTORE, APPLICATIVO INFORMATICO

    Ogni attività conserva il legame con il proprio esecutore e applicativo.
    Gli esecutori e gli applicativi vengono anche aggiunti come entità
    separate (per il matching GUID), ma il legame resta sull'attività.
    """
    entities = []
    executors_seen = set()
    apps_seen = set()

    # Split per blocchi attività — sceglie la regex in base al formato:
    # RTF (striprtf) usa pipe: \n010|TITOLO\n...
    # Antiword usa bell char: 010\x07TITOLO\n...
    if '\x07' in text:
        # Formato antiword (binario .doc) — usa bell char come separatore
        blocks = re.split(r'(\d{2,4})\|?[\x07]?(?=TITOLO)', text)
    else:
        # Formato RTF (striprtf) — usa pipe come separatore
        blocks = re.split(r'\n(\d{2,4})\|', text)

    for i in range(1, len(blocks), 2):
        code = blocks[i]
        content = blocks[i + 1] if i + 1 < len(blocks) else ""

        # Estrai titolo attività
        title_match = re.search(r'TITOLO[\x07]?\s*\n(.+?)(?:\nDESCRIZIONE|\n)', content)
        if not title_match:
            continue

        title = title_match.group(1).strip()

        # Estrai descrizione
        desc_match = re.search(
            r'DESCRIZIONE[\x07]?\s*\n(.+?)(?:\nALTRO STRUMENTO|\nGESTIONE ANOMALIA)',
            content, re.DOTALL
        )
        desc = desc_match.group(1).strip() if desc_match else None

        # Estrai esecutore (nome)
        executor_name = None
        exec_match = re.search(r'ESECUTORE[\x07]?\s*\n(.+)', content)
        if exec_match:
            executor_raw = exec_match.group(1).strip()
            if executor_raw and executor_raw != '-':
                executor_name = executor_raw

        # Estrai applicativo (nome)
        app_name = None
        app_match = re.search(r'APPLICATIVO INFORMATICO[\x07]?\s*\n(.+)', content)
        if app_match:
            app_raw = app_match.group(1).strip()
            if app_raw and app_raw != '-':
                app_name = app_raw

        # Aggiungi l'attività CON le relazioni
        entities.append(WordEntity(
            name=title,
            entity_type='activity',
            code=code,
            description=desc,
            executor=executor_name,
            application=app_name
        ))

        # Aggiungi esecutore come entità separata (per matching GUID)
        if executor_name and executor_name not in executors_seen:
            executors_seen.add(executor_name)
            entities.append(WordEntity(
                name=executor_name,
                entity_type='executor'
            ))

        # Aggiungi applicativo come entità separata (per matching GUID)
        if app_name and app_name not in apps_seen:
            apps_seen.add(app_name)
            entities.append(WordEntity(
                name=app_name,
                entity_type='application'
            ))

    return entities


def summarize_entities(entities: list[WordEntity]) -> dict:
    """Produce un riepilogo delle entità estratte."""
    summary = {}
    for e in entities:
        summary.setdefault(e.entity_type, []).append(e.name)
    return summary


def summarize_relationships(entities: list[WordEntity]) -> list[dict]:
    """
    Produce un riepilogo delle relazioni estratte.
    Restituisce una lista di dizionari con attività, esecutore, applicativo.
    Utile per debug e per il futuro confronto con le connessioni ARIS.
    """
    relationships = []
    for e in entities:
        if e.entity_type == 'activity' and (e.executor or e.application):
            rel = {
                'activity': e.name,
                'code': e.code,
            }
            if e.executor:
                rel['executor'] = e.executor
            if e.application:
                rel['application'] = e.application
            relationships.append(rel)
    return relationships